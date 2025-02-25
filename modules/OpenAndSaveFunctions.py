import base64
import os
from PyQt5.QtCore import Qt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from PyQt5.QtGui import QTextCursor, QTextBlockFormat
from PyQt5.QtWidgets import QTextEdit
from docx import Document
from docx.shared import Pt, RGBColor


def open_docx(file_path, text_edit: QTextEdit):
    temp_dir = os.path.join(os.path.dirname(file_path), "temp_images")
    if not os.path.exists(temp_dir):
        os.makedirs(temp_dir)

    doc = Document(file_path)
    html_content = ""

    def add_styles(run):
        styles = ""
        if run.bold:
            styles += "font-weight:bold;"
        if run.italic:
            styles += "font-style:italic;"
        if run.underline:
            styles += "text-decoration:underline;"
        if run.font.size:
            styles += f"font-size:{run.font.size.pt}pt;"
        if run.font.color and run.font.color.rgb:
            red, green, blue = run.font.color.rgb
            styles += f"color:rgb({red},{green},{blue});"
        return styles

    for para in doc.paragraphs:
        for run in para.runs:
            if run.text:
                styles = add_styles(run)
                html_content += f'<span style="{styles}">{run.text.replace(" ", "&nbsp;")}</span> '
            for inline_shape in run.element.findall(
                    './/{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'):
                for blip in inline_shape.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip'):
                    rID = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                    image_part = doc.part.related_parts[rID]
                    image_data = image_part.blob
                    encoded_image = base64.b64encode(image_data).decode("utf-8")
                    img_tag = f'<img src="data:image/png;base64,{encoded_image}"/>'
                    html_content += img_tag + ' '
        html_content += '<br>'

    text_edit.setHtml(html_content)


def add_image(paragraph, image_path):
    run = paragraph.add_run()
    run.add_picture(image_path)


def save_docx(file_path, text_edit):
    doc = Document()
    cursor = QTextCursor(text_edit.document())

    def apply_char_format(run, char_format):
        if char_format.fontItalic():
            run.italic = True
        if char_format.fontWeight() == 75:
            run.bold = True
        if char_format.fontUnderline():
            run.underline = True
        if char_format.fontPointSize() > 0:
            run.font.size = Pt(char_format.fontPointSize())
        if char_format.foreground().color().isValid():
            qt_color = char_format.foreground().color()
            red = qt_color.red()
            green = qt_color.green()
            blue = qt_color.blue()
            run.font.color.rgb = RGBColor(red, green, blue)

    def apply_block_format(paragraph, block_format):
        # Применение выравнивания
        alignment = block_format.alignment()
        if alignment == Qt.AlignLeft:
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        elif alignment == Qt.AlignCenter:
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        elif alignment == Qt.AlignRight:
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        elif alignment == Qt.AlignJustify:
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        # Применение межстрочных интервалов
        line_height = block_format.lineHeight()
        if block_format.lineHeightType() == QTextBlockFormat.ProportionalHeight:
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            paragraph.paragraph_format.line_spacing = line_height / 100
        elif block_format.lineHeightType() == QTextBlockFormat.FixedHeight:
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            paragraph.paragraph_format.line_spacing = Pt(line_height)

        # Применение отступов
        if block_format.leftMargin() > 0:
            paragraph.paragraph_format.left_indent = Pt(block_format.leftMargin() / 10)
        if block_format.rightMargin() > 0:
            paragraph.paragraph_format.right_indent = Pt(block_format.rightMargin() / 10)
        if block_format.topMargin() > 0:
            paragraph.paragraph_format.space_before = Pt(block_format.topMargin() / 10)
        if block_format.bottomMargin() > 0:
            paragraph.paragraph_format.space_after = Pt(block_format.bottomMargin() / 10)

        # Применение отступа
        if block_format.indent() > 0:
            paragraph.paragraph_format.first_line_indent = Pt(block_format.indent() * 15)

    block = cursor.block()

    while block.isValid():
        paragraph = doc.add_paragraph()
        apply_block_format(paragraph, block.blockFormat())
        iter = block.begin()
        while iter != block.end():
            fragment = iter.fragment()
            if fragment.isValid():
                if fragment.charFormat().isImageFormat():
                    image_format = fragment.charFormat().toImageFormat()
                    image_path = image_format.name()
                    add_image(paragraph, image_path)
                else:
                    run = paragraph.add_run(fragment.text())
                    apply_char_format(run, fragment.charFormat())
            iter += 1
        block = block.next()
    doc.save(file_path)
